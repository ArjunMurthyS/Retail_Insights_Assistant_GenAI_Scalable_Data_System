from docx import Document
from docx.shared import Inches, Pt
import os

# Paths
RESULTS_DIR = r"c:\Users\arjun\OneDrive\Documents\Projects\MyProjects\Retail_Insights_Assistant_GenAI_Scalable_Data_System\results"
IMAGE_PATH = r"c:\Users\arjun\OneDrive\Documents\Projects\MyProjects\Retail_Insights_Assistant_GenAI_Scalable_Data_System\results\retail_insights_assistant_view_1766633553898.png"
OUTPUT_FILE = os.path.join(RESULTS_DIR, "Architecture_Presentation.docx")

document = Document()

# Title
document.add_heading('Retail Insights Assistant', 0)
document.add_heading('GenAI + Scalable Data System', level=1)


# Slide 1: Architecture
document.add_heading('1. System Architecture & Data Flow', level=2)
ARCH_IMAGE_PATH = os.path.join(RESULTS_DIR, "Architecture_Diagram.png")
if os.path.exists(ARCH_IMAGE_PATH):
    document.add_picture(ARCH_IMAGE_PATH, width=Inches(6))
    document.add_paragraph('Figure 1: High-Level System Architecture and Agent Workflow')
else:
    document.add_paragraph('[Architecture Diagram not found]')

p = document.add_paragraph('The system follows a modular Multi-Agent Architecture designed for extensibility and scalability.')
document.add_paragraph('Workflow:', style='List Bullet')
document.add_paragraph('Orchestrator: Manages workflow.', style='List Bullet 2')
document.add_paragraph('Planner: Decomposes complex questions.', style='List Bullet 2')
document.add_paragraph('Data Engineer: Generates safe DuckDB SQL.', style='List Bullet 2')
document.add_paragraph('Analyst: Summarizes results into business insights.', style='List Bullet 2')
document.add_paragraph('Validator: Prevents SQL injection and Hallucinations.', style='List Bullet 2')

# Slide 2: UI Snapshot
document.add_heading('2. System Interface (Snapshot)', level=2)
document.add_paragraph('Below is a live snapshot of the Retail Insights Assistant interface:')
if os.path.exists(IMAGE_PATH):
    document.add_picture(IMAGE_PATH, width=Inches(6))
else:
    document.add_paragraph(f"[Image not found at {IMAGE_PATH}]")

# Slide 3: Scalability
document.add_heading('3. Scalability Strategy (100GB+)', level=2)
document.add_paragraph('A. Data Engineering & Storage', style='Heading 3')
document.add_paragraph('Current: Local DuckDB.')
document.add_paragraph('Proposed (100GB+): Snowflake/BigQuery + S3 Data Lake (Parquet).')

document.add_paragraph('B. Retrieval Efficiency', style='Heading 3')
document.add_paragraph('Partitioning: Date/Region based.')
document.add_paragraph('Vector Search: Pinecone/Qdrant for semantic search.')

# Slide 4: LLM Integration
document.add_heading('4. LLM Integration Strategy', level=2)
document.add_paragraph('Hub-and-Spoke Model:')
document.add_paragraph('Complex Logic: GPT-4o / Gemini 1.5 Pro.')
document.add_paragraph('Task Specific: GPT-3.5 / Llama 3.')
document.add_paragraph('RAG: Retrieve schema/values before prompting.')

document.save(OUTPUT_FILE)
print(f"Presentation saved to {OUTPUT_FILE}")
