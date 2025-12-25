from docx import Document
from docx.shared import Inches, Pt
import os
import glob

# Paths
BASE_DIR = r"c:\Users\arjun\OneDrive\Documents\Projects\MyProjects\Retail_Insights_Assistant_GenAI_Scalable_Data_System"
RESULTS_DIR = os.path.join(BASE_DIR, "results")
README_PATH = os.path.join(RESULTS_DIR, "README.md")
IMAGE_PATH = os.path.join(RESULTS_DIR, "Snapshot_Streamlit_UI.png")
DOCKER_IMAGE_PATH = os.path.join(RESULTS_DIR, "Snapshot_Docker_Run.png")

# --- Generate Technical Notes DOCX ---
def generate_technical_notes():
    doc = Document()
    doc.add_heading('Retail Insights Assistant - Technical Notes', 0)
    
    # Read README content
    if os.path.exists(README_PATH):
        with open(README_PATH, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. '):
                doc.add_paragraph(line[3:], style='List Number')
            elif line:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("README.md not found. Please refer to the source code.")

    output_path = os.path.join(RESULTS_DIR, "Technical_Notes.docx")
    doc.save(output_path)
    print(f"Saved {output_path}")

# --- Generate Demo Evidence DOCX ---
def generate_demo_evidence():
    doc = Document()
    doc.add_heading('Retail Insights Assistant - Demo Evidence', 0)
    
    # 1. Implementation Screenshot
    doc.add_heading('1. User Interface Snapshot (Local)', level=1)
    if os.path.exists(IMAGE_PATH):
        doc.add_picture(IMAGE_PATH, width=Inches(6))
        doc.add_paragraph("Figure 1: Streamlit Chat Interface (Local Run)")
    else:
        doc.add_paragraph(f"[Image not found at {IMAGE_PATH}]")
    
    # 1.5 Docker Implementation Screenshot
    if os.path.exists(DOCKER_IMAGE_PATH):
        doc.add_heading('1.1 Docker Container Run', level=1)
        doc.add_picture(DOCKER_IMAGE_PATH, width=Inches(6))
        doc.add_paragraph("Figure 2: Application running in Docker with 'Summarize' query response.")
    
    # 2. Example Output (Summarization)
    doc.add_heading('2. Example Summarization Output', level=1)
    doc.add_paragraph("Query: 'Summarize sales performance'")
    
    # Find result text files
    result_files = glob.glob(os.path.join(RESULTS_DIR, "result_*.txt"))
    if result_files:
        # Sort by creation time to get the latest
        latest_file = max(result_files, key=os.path.getctime)
        with open(latest_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc.add_paragraph("System Response:", style='Intense Quote')
        doc.add_paragraph(content)
    else:
        doc.add_paragraph("[No result text files found]")

    output_path = os.path.join(RESULTS_DIR, "Demo_Evidence.docx")
    doc.save(output_path)
    print(f"Saved {output_path}")

if __name__ == "__main__":
    try:
        generate_technical_notes()
        generate_demo_evidence()
    except Exception as e:
        print(f"Error generation docs: {e}")
